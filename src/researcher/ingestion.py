"""File ingestion and export.

Upload files (TXT, CSV, PDF, DOCX, XLSX, JSON, MD), extract text content
for the research agent, and export model output to various formats.

Public API
----------
- ``register_ingestion_routes(app)`` — attach ``/files/*`` and ``/export``
- ``init_files_table(db)`` — create the ``files`` table if missing
- ``get_file_context(db, user_id, file_ids)`` — merged text for chat context
"""

import io
import re
import json
import uuid
import logging
from pathlib import Path
from datetime import datetime, timezone

from fastapi import HTTPException, Request, UploadFile, File
from fastapi.responses import Response
from pydantic import BaseModel

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

MAX_FILE_SIZE = 10 * 1024 * 1024        # 10 MB per file
MAX_TEXT_CHARS = 50_000                  # truncate extracted text
MAX_FILES_PER_USER = 50
ALLOWED_EXTENSIONS = {'.txt', '.csv', '.pdf', '.docx', '.xlsx', '.md', '.json'}

UPLOAD_DIR = Path(__file__).parent / "data" / "uploads"

# ---------------------------------------------------------------------------
# DB
# ---------------------------------------------------------------------------


async def init_files_table(db):
    """Create the ``files`` table if it doesn't exist."""
    await db.execute("""
        CREATE TABLE IF NOT EXISTS files (
            id              TEXT PRIMARY KEY,
            user_id         TEXT NOT NULL,
            filename        TEXT NOT NULL,
            extension       TEXT NOT NULL,
            size            INTEGER NOT NULL,
            extracted_text  TEXT NOT NULL DEFAULT '',
            created_at      TEXT NOT NULL
        )
    """)
    await db.commit()


# ---------------------------------------------------------------------------
# Context helper  (called by /chat and /ask)
# ---------------------------------------------------------------------------


async def get_file_context(db, user_id: str, file_ids: list[str]) -> str:
    """Fetch extracted text for *file_ids* and return formatted context block."""
    if not file_ids:
        return ""
    parts = []
    for fid in file_ids:
        cursor = await db.execute(
            "SELECT filename, extracted_text FROM files WHERE id = ? AND user_id = ?",
            (fid, user_id),
        )
        row = await cursor.fetchone()
        if row:
            parts.append(f"--- {row[0]} ---\n{row[1]}")
    if not parts:
        return ""
    return "[Attached Files]\n" + "\n\n".join(parts) + "\n[End of Attached Files]\n\n"


# ---------------------------------------------------------------------------
# Text extractors
# ---------------------------------------------------------------------------


def _decode_bytes(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_txt(data: bytes) -> str:
    return _decode_bytes(data)


def _extract_csv(data: bytes) -> str:
    return _decode_bytes(data)


def _extract_md(data: bytes) -> str:
    return _decode_bytes(data)


def _extract_json(data: bytes) -> str:
    text = _decode_bytes(data)
    try:
        return json.dumps(json.loads(text), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return text


def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF support requires PyMuPDF: pip install pymupdf",
        )
    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text}")
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="DOCX support requires python-docx: pip install python-docx",
        )
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text for cell in row.cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="XLSX support requires openpyxl: pip install openpyxl",
        )
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for name in wb.sheetnames:
        ws = wb[name]
        lines = [f"[Sheet: {name}]"]
        for row in ws.iter_rows(values_only=True):
            lines.append("\t".join(str(c) if c is not None else "" for c in row))
        parts.append("\n".join(lines))
    wb.close()
    return "\n\n".join(parts)


_EXTRACTORS = {
    ".txt": _extract_txt,
    ".csv": _extract_csv,
    ".md": _extract_md,
    ".json": _extract_json,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
}


def _extract_text(data: bytes, ext: str) -> str:
    func = _EXTRACTORS.get(ext)
    if not func:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
    text = func(data)
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + f"\n\n[… truncated at {MAX_TEXT_CHARS:,} characters]"
    return text


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------


def _add_formatted_runs(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add runs to *paragraph*."""
    from docx.shared import Pt

    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def _export_docx(content: str, images_dir: Path | None = None) -> bytes:
    """Convert markdown-ish *content* to DOCX bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX export requires python-docx")

    doc = Document()
    lines = content.split("\n")
    i = 0
    in_code = False
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # ── Code fence ──
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                i += 1
                continue
            in_code = True
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            i += 1
            continue

        # ── Table rows ──
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not all(set(c.strip()) <= set("-: ") for c in cells):
                table_rows.append(cells)
            if not in_table:
                in_table = True
            i += 1
            continue
        elif in_table:
            in_table = False
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, rd in enumerate(table_rows):
                    for ci, ct in enumerate(rd):
                        if ci < ncols:
                            tbl.rows[ri].cells[ci].text = ct
                if table_rows:
                    for cell in tbl.rows[0].cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
            table_rows = []
            continue  # reprocess current line

        # ── Images ──
        img_match = re.match(
            r"!\[([^\]]*)\]\((/static/generated/[^)]+)\)", line.strip()
        )
        if img_match and images_dir:
            img_name = Path(img_match.group(2)).name
            img_full = images_dir / img_name
            if img_full.exists():
                doc.add_picture(str(img_full), width=Inches(5))
            else:
                doc.add_paragraph(f"[Image: {img_match.group(1)}]")
            i += 1
            continue

        # ── Headings ──
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # ── Bullet list ──
        elif re.match(r"^\s*[-*]\s", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, text)
        # ── Numbered list ──
        elif re.match(r"^\s*\d+\.\s", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, text)
        # ── Blank line ──
        elif not line.strip():
            pass
        # ── Math block $$…$$ ──
        elif line.strip().startswith("$$"):
            eq_lines = [line.strip().lstrip("$")]
            i += 1
            while i < len(lines) and "$$" not in lines[i]:
                eq_lines.append(lines[i])
                i += 1
            if i < len(lines):
                eq_lines.append(lines[i].strip().rstrip("$"))
            eq_text = "\n".join(l for l in eq_lines if l.strip())
            p = doc.add_paragraph()
            p.alignment = 1  # center
            run = p.add_run(eq_text)
            run.italic = True
            run.font.size = Pt(11)
        # ── Regular paragraph ──
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)

        i += 1

    # Flush remaining table
    if table_rows:
        ncols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=ncols)
        tbl.style = "Table Grid"
        for ri, rd in enumerate(table_rows):
            for ci, ct in enumerate(rd):
                if ci < ncols:
                    tbl.rows[ri].cells[ci].text = ct

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_xlsx(content: str) -> bytes:
    """Extract markdown tables from *content* and write them to XLSX sheets."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except ImportError:
        raise HTTPException(status_code=500, detail="XLSX export requires openpyxl")

    wb = Workbook()
    wb.remove(wb.active)

    table_pat = re.compile(r"((?:^\|.*\|$\n?)+)", re.MULTILINE)
    tables = table_pat.findall(content)

    if not tables:
        ws = wb.create_sheet("Content")
        for row_idx, line in enumerate(content.split("\n"), 1):
            ws.cell(row=row_idx, column=1, value=line)
    else:
        for tidx, tbl_text in enumerate(tables):
            ws = wb.create_sheet(f"Table {tidx + 1}")
            drow = 0
            for row_text in tbl_text.strip().split("\n"):
                cells = [c.strip() for c in row_text.strip().strip("|").split("|")]
                if all(set(c.strip()) <= set("-: ") for c in cells):
                    continue
                drow += 1
                for ci, ct in enumerate(cells, 1):
                    cell = ws.cell(row=drow, column=ci, value=ct)
                    if drow == 1:
                        cell.font = Font(bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------


class ExportRequest(BaseModel):
    content: str
    format: str            # "md", "txt", "docx", "xlsx"
    filename: str = "export"


# ---------------------------------------------------------------------------
# Route registration
# ---------------------------------------------------------------------------


def register_ingestion_routes(app):
    """Attach ``/files/*`` upload/list/delete and ``/export`` to *app*."""
    from researcher.auth import get_current_user

    # ── Upload ──

    @app.post("/files/upload")
    async def upload_files(
        request: Request,
        files: list[UploadFile] = File(...),
    ):
        user = await get_current_user(request)
        db = request.app.state.db

        cursor = await db.execute(
            "SELECT COUNT(*) FROM files WHERE user_id = ?", (user["id"],)
        )
        count = (await cursor.fetchone())[0]
        if count + len(files) > MAX_FILES_PER_USER:
            raise HTTPException(
                status_code=400,
                detail=f"File limit exceeded (max {MAX_FILES_PER_USER})",
            )

        results = []
        for f in files:
            name = f.filename or "unnamed"
            ext = Path(name).suffix.lower()
            if ext not in ALLOWED_EXTENSIONS:
                results.append({"filename": name, "error": f"Unsupported type: {ext}"})
                continue

            data = await f.read()
            if len(data) > MAX_FILE_SIZE:
                results.append({"filename": name, "error": "File too large (max 10 MB)"})
                continue

            try:
                extracted = _extract_text(data, ext)
            except HTTPException as e:
                results.append({"filename": name, "error": e.detail})
                continue

            fid = uuid.uuid4().hex[:10]
            now = datetime.now(timezone.utc).isoformat()

            # Persist original on disk
            user_dir = UPLOAD_DIR / user["id"]
            user_dir.mkdir(parents=True, exist_ok=True)
            (user_dir / f"{fid}_{name}").write_bytes(data)

            await db.execute(
                "INSERT INTO files "
                "(id, user_id, filename, extension, size, extracted_text, created_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (fid, user["id"], name, ext, len(data), extracted, now),
            )
            await db.commit()

            results.append({
                "id": fid,
                "filename": name,
                "size": len(data),
                "chars": len(extracted),
            })

        return {"files": results}

    # ── List ──

    @app.get("/files")
    async def list_files(request: Request):
        user = await get_current_user(request)
        db = request.app.state.db
        cursor = await db.execute(
            "SELECT id, filename, extension, size, created_at FROM files "
            "WHERE user_id = ? ORDER BY created_at DESC",
            (user["id"],),
        )
        rows = await cursor.fetchall()
        return {
            "files": [
                {
                    "id": r[0], "filename": r[1], "extension": r[2],
                    "size": r[3], "created_at": r[4],
                }
                for r in rows
            ]
        }

    # ── Get one (includes extracted text) ──

    @app.get("/files/{file_id}")
    async def get_file(file_id: str, request: Request):
        user = await get_current_user(request)
        db = request.app.state.db
        cursor = await db.execute(
            "SELECT id, filename, extension, size, extracted_text, created_at "
            "FROM files WHERE id = ? AND user_id = ?",
            (file_id, user["id"]),
        )
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="File not found")
        return {
            "id": row[0], "filename": row[1], "extension": row[2],
            "size": row[3], "text": row[4], "created_at": row[5],
        }

    # ── Delete ──

    @app.delete("/files/{file_id}")
    async def delete_file(file_id: str, request: Request):
        user = await get_current_user(request)
        db = request.app.state.db
        cursor = await db.execute(
            "SELECT id FROM files WHERE id = ? AND user_id = ?",
            (file_id, user["id"]),
        )
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="File not found")

        # Remove from disk
        user_dir = UPLOAD_DIR / user["id"]
        for p in user_dir.glob(f"{file_id}_*"):
            p.unlink(missing_ok=True)

        await db.execute("DELETE FROM files WHERE id = ?", (file_id,))
        await db.commit()
        return {"deleted": file_id}

    # ── Export ──

    @app.post("/export")
    async def export_content(req: ExportRequest, request: Request):
        await get_current_user(request)
        fmt = req.format.lower().strip()
        safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", req.filename)[:50] or "export"

        if fmt == "md":
            return Response(
                content=req.content.encode(),
                media_type="text/markdown",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}.md"'},
            )

        if fmt == "txt":
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", req.content)
            text = re.sub(r"\*([^*]+)\*", r"\1", text)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
            return Response(
                content=text.encode(),
                media_type="text/plain",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}.txt"'},
            )

        if fmt == "docx":
            images_dir = Path(__file__).parent / "static" / "generated"
            data = _export_docx(req.content, images_dir)
            return Response(
                content=data,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
            )

        if fmt == "xlsx":
            data = _export_xlsx(req.content)
            return Response(
                content=data,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}.xlsx"'},
            )

        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")
